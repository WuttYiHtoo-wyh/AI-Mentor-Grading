import os
import re
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document


DOCUMENT_TYPE_MAP = {
    'assignmentbrief': 'assignment_brief',
    'assessmentrubric': 'rubric',
    'goodlearnerreport': 'sample_good',
    'badlearnerreport': 'sample_bad',
    'required_doc': 'required_document',
}

METADATA_FIELD_MAP = {
    'document type': 'document_type',
    'module': 'module',
    'instructional unit': 'instructional_unit',
    'instructional_unit': 'instructional_unit',
    'section': 'section',
    'topic': 'topic',
    'version': 'version',
    'source': 'source',
    'audience': 'audience',
    'intent': 'intent',
    'difficulty': 'difficulty',
    'chunk id': 'chunk_id',
    'title': 'title',
    'learning objective': 'learning_objective',
    'content': 'content',
    'example': 'example',
    'key takeaways': 'key_takeaways',
    'question': 'question',
    'keywords': 'keywords',
}


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return ' '.join(lines)


def determine_document_type(filename: str) -> str:
    key = Path(filename).stem.lower().replace(' ', '').replace('-', '').replace('_', '')
    for token, doc_type in DOCUMENT_TYPE_MAP.items():
        if token in key:
            return doc_type
    if key.startswith('referencematerial'):
        return 'reference_material'
    return 'unknown'


def normalize_metadata_key(key: str) -> str:
    normalized = key.strip().lower().replace(' ', ' ')
    return METADATA_FIELD_MAP.get(normalized, normalized.replace(' ', '_'))


def parse_metadata_string(raw_metadata: str) -> Dict[str, object]:
    metadata: Dict[str, object] = {}
    lines = [line.strip() for line in re.split(r'\r?\n|\u2028|\u2029', raw_metadata) if line.strip()]
    for line in lines:
        if ':' in line:
            key, value = line.split(':', 1)
            normalized_key = normalize_metadata_key(key)
            value = value.strip()
            if normalized_key == 'keywords':
                keywords = [kw.strip() for kw in re.split(r'[;,\n]+', value) if kw.strip()]
                metadata[normalized_key] = ', '.join(keywords)
            else:
                metadata[normalized_key] = value
        else:
            if 'notes' in metadata:
                metadata['notes'] = f"{metadata['notes']} {line}"
            else:
                metadata['notes'] = line
    return metadata


def parse_docx_chunk_tables(document: Document) -> List[Dict[str, Optional[str]]]:
    chunks: List[Dict[str, Optional[str]]] = []
    for table in document.tables:
        first_cell = table.rows[0].cells[0].text.strip().lower() if table.rows and table.rows[0].cells else ''
        if 'chunk id' not in first_cell:
            continue

        chunk: Dict[str, Optional[str]] = {}
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 1 or not cells[0]:
                continue
            field_key = normalize_metadata_key(cells[0])
            value = cells[1] if len(cells) > 1 else ''
            if field_key == 'metadata':
                chunk.update(parse_metadata_string(value))
            elif field_key == 'keywords':
                keywords = [kw.strip() for kw in re.split(r'[;,\n]+', value) if kw.strip()]
                chunk[field_key] = ', '.join(keywords)
            else:
                chunk[field_key] = normalize_text(value)

        if chunk.get('chunk_id') or chunk.get('title') or chunk.get('content'):
            text_parts = [chunk.get('title'), chunk.get('learning_objective'), chunk.get('content'), chunk.get('example'), chunk.get('key_takeaways')]
            normalized_text = normalize_text(' '.join([part for part in text_parts if part]))
            chunk['text'] = normalized_text
            chunks.append(chunk)
    return chunks


def extract_docx_text(filepath: Path) -> List[Dict[str, Optional[str]]]:
    if not filepath.exists():
        raise FileNotFoundError(f'Document not found: {filepath}')

    try:
        document = Document(filepath)
    except Exception as exc:
        raise ValueError(f'Failed to open Word document {filepath}: {exc}')

    chunk_tables = parse_docx_chunk_tables(document)
    if chunk_tables:
        return chunk_tables

    chunks: List[Dict[str, Optional[str]]] = []
    current_heading = None
    current_text: List[str] = []

    def flush_paragraphs():
        nonlocal current_heading, current_text
        if not current_text:
            return
        text = normalize_text(' '.join(current_text))
        if text:
            chunks.append({'heading': current_heading, 'text': text})
        current_text = []

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name.lower() if paragraph.style else ''
        text = paragraph.text.strip()
        if not text:
            continue

        if 'heading' in style_name or 'title' in style_name or 'subtitle' in style_name:
            flush_paragraphs()
            current_heading = text
            continue

        current_text.append(text)

    flush_paragraphs()

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(' | '.join(cells))
        if rows:
            table_text = normalize_text(' '.join(rows))
            chunks.append({'heading': f'Table {table_index}', 'text': table_text})

    return chunks


def chunk_text(text: str, chunk_size: int = 700, overlap: int = 100) -> List[str]:
    if chunk_size <= overlap:
        raise ValueError('chunk_size must be greater than overlap')

    chunks = []
    start = 0
    text_length = len(text)
    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunks.append(text[start:end])
        if end == text_length:
            break
        start += chunk_size - overlap
    return chunks
